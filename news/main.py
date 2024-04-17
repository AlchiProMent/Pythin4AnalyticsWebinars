# Получение текстов статей новостных сайтов
# репозиторий: https://github.com/AlchiProMent/Pythin4AnalyticsWebinars
# email: alchiprotoday@gmail.com

from urllib.request import urlopen, Request
from urllib import error as err

from os import remove
from os.path import exists
import docx
from deep_translator import GoogleTranslator

from const import site_type, _TMP_JPG, _KEY_TITLE, _KEY_AUTOR, _KEY_DATE, _KEY_BODY, _FLUKY_SITE
import const
from parsers import *

# заголовок фиктивного браузера
HEADERS = {'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_11_5) AppleWebKit/537.36 '
                         '(KHTML, like Gecko) Chrome/50.0.2661.102 Safari/537.36'}

# папка для хранения загруженнных и переведенных статей
_ART_PATH = 'TEXTES/'
# ширина изображения (см), вкючаемого в документ
_IMAGE_WIDTH_CM= 14

def get_art_txt(url):
    # получить текст статьи по её URL

    txt = None
    web_type = site_type(url)
    if web_type == _FLUKY_SITE:
        err_msg('Эта страница не может быть правильно распознана')
    else:
        # загрузка статьи по URL, парсинг в DOM и вызов парсера для конкретного сайта
        print('Парсинг страницы...')
        try:

            # создать подключение, передав заголовок фиктивного браузера
            req = Request(url, headers=HEADERS)
            with urlopen(req) as response:
                html = str(response.read(), 'UTF-8')

        except err.HTTPError:
            err_msg('Страница не найдена!')
        except err.URLError:
            err_msg('Неверный URL!')
        except ValueError:
            err_msg('Неизвестный тип URL!')
        else:
            msg('Текст страницы получен')
            dom = BeautifulSoup(html, 'html.parser')

            match web_type:
                case const._NY_POST:
                    # парсить статью из New York Post
                    txt = nypost_art_parser(dom)
                case const._NY_TIMES:
                    txt = nytimes_art_parser(url)

    return txt


def print_art_text(full_body:dict):
    # вывод текста ранее спарсенного текстат статьи на экран

    for key, item in full_body.items():
        if key != _KEY_BODY:
            print(f'{key}:', end=' ')
            prntxt(item)

    body_text = full_body[_KEY_BODY]
    print()
    for el in body_text:
        if len(el.keys()) == 1:
            # обычный абзац или подзаголовок
            for key, text in el.items():
                if key == 'h2':
                    print()
                    print(text.upper())
                else:
                    prntxt(text)
        else:
            # ссылка на картинку
            print(el['src'])
            # подпись
            print(el['figure'])
            if 'srcset' in el:
                for nextsrc in el['srcset']:
                    print(f'{nextsrc[0]} (ширина {nextsrc[1]} px)')

        print()

def save_art_to_word(art_body, url, inc_image=True, translate=False):
    # сохранить ранее спарсенную статью в Word-файл

    # сохранить в Word-файл
    doc = docx.Document()
    # создать переводчик
    ts = GoogleTranslator(source='auto', target='ru')

    title = art_body[_KEY_TITLE]
    file_name = f'{title[:80]}'
    file_name = file_name.replace(':', '')
    file_name = file_name.replace('/', '')
    file_name = file_name.replace('?', '')
    if translate:
        title = ts.translate(title)
        file_name = f'{file_name} - RU'
    file_name = f'{_ART_PATH}{file_name}.docx'
    jpeg_file_name = f'{_ART_PATH}{_TMP_JPG}'

    if _KEY_AUTOR in art_body:
        autor_name = art_body[_KEY_AUTOR]
    else:
        autor_name = None
    pub_date = art_body[_KEY_DATE]
    body = art_body[_KEY_BODY]

    doc.add_heading(title, 0)
    if autor_name:
        doc.add_paragraph(f'Автор: {autor_name}')
    doc.add_paragraph(f'Дата: {pub_date}')
    doc.add_paragraph('')

    for el in body:
        if len(el.keys()) == 1:
            # обычный абзац или подзаголовок
            for key, text in el.items():
                s = ts.translate(text) if translate else text
                if key == 'h2':
                    doc.add_heading(s, 1)
                elif key == 'h3':
                    doc.add_heading(s, 2)
                    doc.add_paragraph('')
                elif key == 'h4':
                    doc.add_heading(s, 3)
                    doc.add_paragraph('')
                elif key == 'h5':
                    doc.add_heading(s, 4)
                    doc.add_paragraph('')
                elif key == 'h6':
                    doc.add_heading(s, 5)
                    doc.add_paragraph('')
                else:
                    doc.add_paragraph(s)
        else:
            img_src = el['src']
            img_txt = el['figure']

            # включить изображение в документ
            if inc_image:
                try:
                    req = Request(img_src, headers=HEADERS)
                    with urlopen(req) as response:
                        jpeg = response.read()

                except:
                    print(f'Не удалось получить изображение по адресу {img_src}!')
                else:
                    out = open(jpeg_file_name, 'wb')
                    out.write(jpeg)
                    out.close()

                    # включить изображение в документ
                    doc.add_picture(jpeg_file_name, width = docx.shared.Cm(_IMAGE_WIDTH_CM))
                    s = ts.translate(img_txt) if translate else img_txt
                    doc.add_heading(s, 6)
                    doc.add_paragraph('')

        # удалить временный файл
        if exists(jpeg_file_name):
            try:
                remove(jpeg_file_name)
            except:
                pass

    doc.add_paragraph('')
    doc.add_paragraph(f'Источник: {url}')

    try:
        doc.save(file_name)
    except Exception as e:
        err_msg(f'Ошибка при сохранении файла "{file_name}"!')
        print(e)
    else:
        msg(f'Статья сохранена в файл "{file_name}"')


if __name__ == '__main__':

    out_on_console = False

    while (article_url := input('\nВведите URL статьи (ENTER для выхода): ')) != '':

        if article_url != '':

            # если пользователь ввел URL статьи - получить текст статьи
            article_txt = get_art_txt(article_url)
            print()

            if article_txt:
                if out_on_console:
                    # вывести на экран
                    print_art_text(article_txt)

                # сохранитиь полученный оригинальный текст в Word-файл
                print('Сохранение оригинального текста...')
                save_art_to_word(article_txt, article_url)
                print('Перевод статьи...')
                save_art_to_word(article_txt, article_url, translate=True)

    print('Программа закончила работу')
